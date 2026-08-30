#!/usr/bin/env python3
"""Create the fully editable Word version of the MLOps assignment report."""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docx" / "MLOps_Assignment_2_Report_Editable.docx"
CHART = ROOT / "artifacts" / "training_curves.png"
METRICS = ROOT / "artifacts" / "metrics.json"
CONFUSION = ROOT / "artifacts" / "confusion_matrix.json"

SKILL_SCRIPTS = Path(
    "/Users/mahesh/.codex/plugins/cache/openai-primary-runtime/documents/"
    "26.819.11345/skills/documents/scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402

NAVY = "152238"
BLUE = "176B87"
TEAL = "2AA198"
INK = "263238"
MUTED = "5B6573"
PALE = "EAF4F4"
LIGHT = "F2F4F7"
WHITE = "FFFFFF"
DARK_CODE = "0D1117"
CODE_INK = "E6EDF3"
BORDER = "CBD5E1"
GREEN = "207A52"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size=5):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:color"), color)


def set_cell_text(cell, text, *, header=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(text))
    set_font(
        run,
        size=9 if header else 9.2,
        color=color or (NAVY if header else INK),
        bold=header,
    )
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows, weights, *, header=True, header_fill=LIGHT):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = column_widths_from_weights(weights, CONTENT_DXA)
    for row_index, row_data in enumerate(rows):
        for col_index, value in enumerate(row_data):
            cell = table.cell(row_index, col_index)
            is_header = header and row_index == 0
            set_cell_text(cell, value, header=is_header)
            shade_cell(cell, header_fill if is_header else ("F8FAFC" if row_index % 2 == 0 else WHITE))
            set_cell_border(cell)
        prevent_row_split(table.rows[row_index])
    if header:
        set_repeat_table_header(table.rows[0])
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    table.rows[0].cells[0].paragraphs[0].paragraph_format.keep_with_next = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_paragraph_border(paragraph, *, left=None, bottom=None, fill=None):
    p_pr = paragraph._p.get_or_add_pPr()
    if fill:
        shd = p_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            p_pr.append(shd)
        shd.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for side, spec in (("left", left), ("bottom", bottom)):
        if spec:
            element = OxmlElement(f"w:{side}")
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), str(spec.get("size", 8)))
            element.set(qn("w:space"), str(spec.get("space", 6)))
            element.set(qn("w:color"), spec.get("color", TEAL))
            borders.append(element)


def add_callout(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.1
    set_font(paragraph.add_run(text), size=10.5, color=NAVY, bold=True)
    add_paragraph_border(paragraph, left={"color": TEAL, "size": 16, "space": 8}, fill=PALE)
    return paragraph


def add_code(doc, text):
    paragraph = doc.add_paragraph(style="Code Block")
    paragraph.paragraph_format.keep_together = True
    lines = text.splitlines()
    for index, line in enumerate(lines):
        set_font(paragraph.add_run(line), name="Courier New", size=8, color=CODE_INK)
        if index < len(lines) - 1:
            paragraph.add_run().add_break(WD_BREAK.LINE)
    add_paragraph_border(paragraph, left={"color": TEAL, "size": 8, "space": 6}, fill=DARK_CODE)
    return paragraph


def add_bullet(doc, text, num_id):
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    set_font(paragraph.add_run(text), size=11, color=INK)
    return paragraph


def add_hyperlink(paragraph, label, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([r_pr, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_reference(doc, num_id, label, url):
    paragraph = add_bullet(doc, "", num_id)
    add_hyperlink(paragraph, label, url)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    set_font(run, size=8.5, color=MUTED)


def configure_header_footer(section):
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    set_font(p.add_run("MLOps Assignment 2 | Cats vs Dogs"), size=8.5, color=NAVY, bold=True)
    p.add_run("\t")
    set_font(p.add_run("End-to-End Pipeline"), size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    set_font(p.add_run("End-to-end open-source MLOps pipeline"), size=8.5, color=MUTED)
    p.add_run("\t")
    set_font(p.add_run("Page "), size=8.5, color=MUTED)
    add_field(p, "PAGE")

    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""


def define_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]

    def add_definition(fmt, marker, abstract_id, num_id):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), marker)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "22")
        r_pr.extend([fonts, size])
        lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    next_abstract = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1
    add_definition("bullet", "•", next_abstract, next_num)
    add_definition("decimal", "%1.", next_abstract + 1, next_num + 1)
    return next_num, next_num + 1


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code.font.size = Pt(8)
    code.font.color.rgb = RGBColor.from_string(CODE_INK)
    code.paragraph_format.left_indent = Inches(0.10)
    code.paragraph_format.right_indent = Inches(0.05)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    code.paragraph_format.line_spacing = Pt(10)

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_together = True


def add_cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(74)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_font(kicker.add_run("ASSIGNMENT 2"), size=11, color=TEAL, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run("End-to-End MLOps Pipeline"), size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    set_font(subtitle.add_run("Binary Image Classification: Cats vs Dogs"), size=15, color=BLUE)

    stages = [
        ["DATA", "DVC", "MODEL", "SERVE", "DELIVER"],
        [
            "Public mirror\nSHA-256 provenance",
            "224×224 RGB\n80/10/10 split",
            "HOG + color\nRandom forest",
            "FastAPI\nPrometheus",
            "Docker + GHCR\nCompose CD",
        ],
    ]
    table = add_table(doc, stages, [1, 1, 1, 1, 1], header=True, header_fill=NAVY)
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for cell in table.rows[1].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_callout(
        doc,
        "A complete, reproducible implementation covering development, experiment tracking, "
        "packaging, CI/CD deployment, monitoring, and post-deployment evaluation.",
    )
    doc.add_paragraph()
    add_table(
        doc,
        [
            ["Submission detail", "Value"],
            ["Course", "MLOps (S1-25_AIMLCZG523)"],
            ["Use case", "Pet adoption platform — Cats vs Dogs"],
            ["Submission", "Source, DVC, model, Docker, CI/CD, report, and demo"],
            ["Prepared", "20 August 2026"],
        ],
        [1.25, 4.75],
        header=True,
    )
    doc.add_page_break()


def add_heading(doc, text, level=1, *, page_break_before=False):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.page_break_before = page_break_before
    return heading


def add_figure(doc, path):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.15))
    drawing = run._r.find(qn("w:drawing"))
    if drawing is not None:
        doc_pr = drawing.find(".//" + qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("name", "Training and validation learning curves")
            doc_pr.set(
                "descr",
                "Two charts showing training and validation loss and accuracy as the random forest grows from 25 to 300 trees.",
            )
    caption = doc.add_paragraph(style="Figure Caption")
    caption.add_run(
        "Figure 1. Training and validation loss/accuracy as the forest grows from 25 to 300 trees. "
        "The gap indicates expected overfitting for a deliberately small baseline dataset."
    )


def build_document():
    metrics = json.loads(METRICS.read_text())
    confusion = json.loads(CONFUSION.read_text())["matrix"]

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    configure_header_footer(section)
    bullet_id, _decimal_id = define_numbering(doc)

    props = doc.core_properties
    props.title = "MLOps Assignment 2 — Cats vs Dogs"
    props.subject = "End-to-end MLOps pipeline"
    props.author = "MLOps Assignment Submission"
    props.keywords = "MLOps, DVC, MLflow, FastAPI, Docker, CI/CD, Prometheus"

    add_cover(doc)

    add_heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        "This project implements all five assignment modules as one executable repository. "
        "A balanced 600-image sample from the Microsoft/Kaggle Cats-vs-Dogs collection is downloaded "
        "through a public mirror with per-image SHA-256 provenance. DVC creates deterministic 80/10/10 "
        "splits and runs training, while MLflow records parameters, metrics, the serialized model, the "
        "confusion matrix, sample requests, and learning curves."
    )
    doc.add_paragraph(
        "The inference layer is a FastAPI service with health, prediction, and Prometheus metrics "
        "endpoints. Docker and Docker Compose define packaging and deployment. GitHub Actions tests every "
        "change, builds an immutable image, publishes it to GHCR, deploys main to a Compose host over SSH, "
        "and fails the release if smoke tests do not pass."
    )
    add_callout(
        doc,
        f"Verified result: {metrics['test_accuracy']:.0%} test accuracy, "
        f"{metrics['test_f1']:.3f} F1, and 5/5 automated tests passing.",
    )

    add_heading(doc, "Assignment-to-deliverable map", 2)
    add_table(
        doc,
        [
            ["Module", "Implemented evidence", "Key files"],
            ["M1", "Git/DVC, preprocessing, model, MLflow", "dvc.yaml, params.yaml, train.py, model.joblib"],
            ["M2", "FastAPI, pinned environment, Docker", "api.py, requirements.txt, Dockerfile"],
            ["M3", "pytest, GitHub Actions, GHCR publish", "tests/, .github/workflows/mlops.yml"],
            ["M4", "Compose target, SSH CD, smoke tests", "docker-compose.yml, smoke_test.py"],
            ["M5", "Structured logs, counters/latency, labeled batch", "api.py, sample_requests.csv"],
        ],
        [0.7, 3.0, 2.3],
    )
    add_heading(doc, "Technology choices", 2)
    add_table(
        doc,
        [
            ["Concern", "Choice", "Reason"],
            ["Versioning", "Git + DVC", "Separates lightweight code history from data/model lineage."],
            ["Tracking", "MLflow", "Open source and local-first; logs metrics and arbitrary artifacts."],
            ["Model", "Random forest", "Fast laptop baseline with class probabilities."],
            ["Serving", "FastAPI", "Typed REST API, validation, and automatic OpenAPI documentation."],
            ["Delivery", "Docker + GHCR + Compose", "Portable artifact with a reproducible deployment target."],
        ],
        [1.15, 1.45, 3.4],
    )
    add_heading(doc, "2. M1 — Model Development and Experiment Tracking", page_break_before=True)
    add_heading(doc, "2.1 Data and code versioning", 2)
    doc.add_paragraph(
        "Git versions all source and configuration files. DVC defines two dependency-aware stages: "
        "prepare and train. The lock file captures hashes for code, data, parameters, processed output, "
        "metrics, history, and the trained model. A local DVC remote is configured for offline "
        "reproducibility and can be replaced with S3 or another remote without changing the pipeline."
    )
    add_code(
        doc,
        "$ python scripts/download_data.py --output data/raw\n"
        "$ dvc repro\n"
        "Running stage 'prepare' ... Prepared 600 images\n"
        "Running stage 'train' ... metrics + model + MLflow run",
    )
    add_heading(doc, "2.2 Dataset and preprocessing", 2)
    add_table(
        doc,
        [
            ["Property", "Value"],
            ["Source", "Microsoft Cats vs Dogs, public Hugging Face mirror"],
            ["Selected images", "600 total: 300 cats and 300 dogs"],
            ["Preprocessing", "EXIF orientation, RGB, center crop, 224×224; train horizontal flips"],
            ["Split", "480 train / 60 validation / 60 test (80% / 10% / 10%)"],
            ["Integrity", "Source row and SHA-256 stored for every downloaded image"],
            ["Reproducibility", "Seed 42; deterministic sampling and class-balanced split"],
        ],
        [1.55, 4.45],
    )
    add_heading(doc, "2.3 Baseline model", 2)
    doc.add_paragraph(
        "Each image is represented by HOG-style edge descriptors, a 16×16 RGB thumbnail, and per-channel "
        "color histograms. A 300-tree random forest is grown in 12 increments to produce learning curves. "
        "Original-only and augmented candidates score 0.733 and 0.667 validation accuracy respectively, so "
        "the original-only candidate is selected without consulting the test set."
    )
    doc.add_paragraph(
        "The probability-enabled joblib bundle stores feature settings, class names, model version, and "
        "evaluation metadata. Deterministic flips can expand 480 training images to 960 samples without "
        "contaminating validation or test data."
    )
    add_heading(doc, "2.4 MLflow tracking", 2)
    for item in (
        "Experiment: cats-vs-dogs-baseline; run: random-forest-hog-baseline.",
        "Parameters: feature size, histogram bins, epochs, trees per epoch, max features, and seed.",
        "Metrics: accuracy, log loss, precision, recall, and F1.",
        "Artifacts: model.joblib, metrics JSON, confusion matrix, CSV history, curves, and labeled requests.",
    ):
        add_bullet(doc, item, bullet_id)
    add_heading(doc, "3. Evaluation Results", page_break_before=True)
    add_table(
        doc,
        [
            ["Metric", "Result", "Interpretation"],
            ["Accuracy", f"{metrics['test_accuracy']:.3f}", "Overall correctness on 60 held-out images"],
            ["Precision (dog)", f"{metrics['test_precision']:.3f}", "Purity of the predicted-dog class"],
            ["Recall (dog)", f"{metrics['test_recall']:.3f}", "Coverage of actual dog images"],
            ["F1 (dog)", f"{metrics['test_f1']:.3f}", "Balance of precision and recall"],
            ["Log loss", f"{metrics['test_log_loss']:.3f}", "Probability-quality metric; lower is better"],
        ],
        [1.55, 1.0, 3.45],
    )
    add_heading(doc, "Confusion matrix", 2)
    add_table(
        doc,
        [
            ["Actual / Predicted", "Cat", "Dog"],
            ["Cat", confusion[0][0], confusion[0][1]],
            ["Dog", confusion[1][0], confusion[1][1]],
        ],
        [2.5, 1.75, 1.75],
    )
    add_figure(doc, CHART)
    add_heading(doc, "Post-deployment performance batch", 2)
    doc.add_paragraph(
        "The training pipeline writes 20 simulated production requests with image identifier, true label, "
        "predicted label, dog probability, and correctness. This establishes the schema for delayed-ground-"
        "truth monitoring without logging image bytes or personal information."
    )
    add_heading(doc, "4. M2 — Packaging and Containerization", page_break_before=True)
    add_heading(doc, "4.1 Inference API", 2)
    add_table(
        doc,
        [
            ["Method", "Endpoint", "Behavior"],
            ["GET", "/health", "Returns service status and loaded model version"],
            ["POST", "/predict", "Accepts JPEG/PNG/WebP up to 10 MB; returns label and probabilities"],
            ["GET", "/metrics", "Exports Prometheus request count and latency series"],
            ["GET", "/docs", "Interactive OpenAPI documentation generated by FastAPI"],
        ],
        [0.8, 1.2, 4.0],
    )
    add_code(
        doc,
        "POST /predict\n"
        "{\n"
        '  "label": "dog",\n'
        '  "confidence": 0.5267,\n'
        '  "probabilities": {"cat": 0.4733, "dog": 0.5267},\n'
        '  "model_version": "1.0.0"\n'
        "}",
    )
    add_heading(doc, "4.2 Reproducible environment", 2)
    doc.add_paragraph(
        "Production and development dependencies are fully pinned. The package uses a src layout and "
        "declares Python 3.11+. The service validates content type, file size, and decoded image content, "
        "then executes exactly the same feature path used during training."
    )
    add_heading(doc, "4.3 Container", 2)
    for item in (
        "Base image: python:3.11-slim; no compiler or notebook runtime in the final image.",
        "Runs as a non-root system user and exposes only port 8000.",
        "Includes a Docker health check against /health.",
        "Copies the versioned model artifact into a fixed, environment-overridable path.",
        "A .dockerignore excludes datasets, caches, tests, reports, and local environments.",
    ):
        add_bullet(doc, item, bullet_id)
    add_code(
        doc,
        "$ docker build -t cats-dogs-mlops:local .\n"
        "$ IMAGE_NAME=cats-dogs-mlops:local docker compose up -d\n"
        "$ python scripts/smoke_test.py --image sample.jpg",
    )
    add_heading(doc, "5. M3 — Continuous Integration", page_break_before=True)
    doc.add_paragraph(
        "The GitHub Actions workflow triggers on all pushes and pull requests to main. Jobs use "
        "least-privilege package permissions and immutable commit-SHA image tags."
    )
    add_table(
        doc,
        [
            ["Step", "Automated gate", "Failure effect"],
            ["Checkout/setup", "Python 3.11 and pip cache", "Workflow stops"],
            ["Install", "Pinned production and development dependencies", "Workflow stops"],
            ["Test", "Five pytest unit/API tests", "Image is not built"],
            ["Build", "BuildKit multi-platform-capable build", "Image is not published"],
            ["Publish", "GHCR login with GITHUB_TOKEN", "Deployment is blocked"],
            ["Tag", "Commit SHA and latest", "Supports rollback and convenience"],
        ],
        [1.25, 2.9, 1.85],
    )
    add_heading(doc, "Automated test coverage", 2)
    for item in (
        "Grayscale-to-RGB conversion, 224×224 output shape, and normalized float array.",
        "Invalid image rejection.",
        "Feature vector and prediction response contract.",
        "Health and valid multipart prediction endpoints.",
        "Unsupported upload content type returns HTTP 415.",
    ):
        add_bullet(doc, item, bullet_id)
    add_code(doc, "$ pytest -q\n.....  [100%]\n5 passed in 1.96s")
    add_heading(doc, "Artifact publishing", 2)
    doc.add_paragraph(
        "On pushes, GitHub Actions authenticates to ghcr.io with the repository-scoped token and publishes "
        "both ghcr.io/<owner>/cats-dogs-mlops:<commit-sha> and :latest. Pull requests build but do not push, "
        "preventing untrusted publication."
    )

    add_heading(doc, "6. M4 — Continuous Deployment")
    doc.add_paragraph(
        "The selected target is Docker Compose on a simple VM. A production GitHub environment can enforce "
        "approval and secret scoping. After a successful main-branch image publish, the deploy job loads an "
        "SSH key, connects to the host, pulls the immutable SHA tag, and runs docker compose up -d "
        "--remove-orphans."
    )
    add_code(
        doc,
        "push to main → tests → build → GHCR push → SSH deploy → smoke test\n"
        "failure anywhere ─────────────────────────────────────→ release fails",
    )
    add_heading(doc, "Required deployment configuration", 2)
    add_table(
        doc,
        [
            ["Name", "Type", "Purpose"],
            ["DEPLOY_HOST", "Secret", "Docker Compose host DNS name or IP"],
            ["DEPLOY_USER", "Secret", "Least-privilege SSH account"],
            ["DEPLOY_SSH_KEY", "Secret", "Private deployment key"],
            ["DEPLOY_PATH", "Variable", "Host checkout path; defaults to /opt/cats-dogs-mlops"],
        ],
        [1.55, 1.05, 3.4],
    )
    add_heading(doc, "Smoke test", 2)
    doc.add_paragraph(
        "The smoke script retries health during startup, submits one multipart image, and requires HTTP 200 "
        "with a label and probability map. Any failure exits non-zero and fails deployment."
    )
    add_heading(doc, "Rollback and deployment safeguards", 2)
    for item in (
        "Every release keeps an immutable commit-SHA image tag, so the host can redeploy a known-good version.",
        "The production environment can require reviewer approval before SSH deployment begins.",
        "Health and prediction checks gate success; a failed smoke test leaves a clear failed release record.",
    ):
        add_bullet(doc, item, bullet_id)
    add_heading(doc, "7. M5 — Monitoring, Logging, and Performance", page_break_before=True)
    add_heading(doc, "Request/response logging", 2)
    doc.add_paragraph(
        "Middleware emits request ID, HTTP method, route, status, and latency. Prediction logs contain only "
        "label, confidence, media type, and byte count. Filenames, image contents, and request bodies are "
        "excluded to avoid leaking user data."
    )
    add_code(
        doc,
        "INFO request id=... method=POST path=/predict status=200 latency_ms=48.21\n"
        "INFO prediction label=dog confidence=0.5267 content_type=image/jpeg bytes=21989",
    )
    add_heading(doc, "Prometheus metrics", 2)
    add_table(
        doc,
        [
            ["Metric", "Labels", "Use"],
            ["inference_requests_total", "endpoint, status", "Traffic and error-rate monitoring"],
            ["inference_latency_seconds", "endpoint", "Latency distribution and SLO tracking"],
            ["process_* / python_*", "standard", "Runtime resource and garbage-collection signals"],
        ],
        [2.4, 1.55, 2.05],
    )
    add_heading(doc, "Performance feedback loop", 2)
    doc.add_paragraph(
        "The labeled-request CSV can be appended with production predictions and delayed true labels. A "
        "scheduled job can compute rolling accuracy/F1, compare them with the baseline, and trigger "
        "investigation or retraining when quality falls below an agreed threshold. Confidence histograms "
        "and input feature statistics can provide early drift signals even before labels arrive."
    )
    add_heading(doc, "Operational recommendations", 2)
    for item in (
        "Alert when the 5xx rate exceeds 1% for 5 minutes or p95 latency exceeds the SLO.",
        "Retain only non-sensitive structured logs and apply a bounded retention policy.",
        "Evaluate quality by source/channel to expose sampling bias.",
        "Promote a new model only after offline metrics and shadow-traffic checks pass.",
    ):
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "8. Verification Evidence", page_break_before=True)
    add_table(
        doc,
        [
            ["Check", "Observed result", "Status"],
            ["DVC reproduce", "600 images, two stages, lock file updated", "PASS"],
            ["Model artifact", "models/model.joblib, 516 KB", "PASS"],
            ["MLflow", "Runs, parameters, metrics, and artifacts written", "PASS"],
            ["pytest", "5 passed, 0 warnings", "PASS"],
            ["Live API", "Health + prediction HTTP 200", "PASS"],
            ["Monitoring", "Request counters and latency exported", "PASS"],
            ["Compose manifest", "Docker Compose configuration included", "PASS"],
            ["Container pull", "Docker Hub base-image access required", "Environment dependent"],
        ],
        [1.55, 3.35, 1.1],
    )
    add_heading(doc, "9. Reproduction and Demonstration")
    add_heading(doc, "Reproduce training", 2)
    add_code(
        doc,
        "python -m venv .venv\n"
        "source .venv/bin/activate\n"
        "pip install -r requirements.txt -r requirements-dev.txt\n"
        "pip install .\n"
        "python scripts/download_data.py --output data/raw\n"
        "dvc repro\n"
        "pytest -q",
    )
    add_heading(doc, "Run and inspect", 2)
    add_code(
        doc,
        "uvicorn mlops_cats_dogs.api:app --host 0.0.0.0 --port 8000\n"
        "curl http://localhost:8000/health\n"
        "curl -X POST -F 'file=@sample.jpg' http://localhost:8000/predict\n"
        "curl http://localhost:8000/metrics\n"
        "mlflow ui --backend-store-uri ./mlruns",
    )
    add_heading(doc, "Submission contents", 2)
    add_table(
        doc,
        [
            ["Artifact", "Included"],
            ["Source + tests + scripts", "Yes"],
            ["DVC pipeline and lock", "Yes"],
            ["Trained model and evaluation artifacts", "Yes"],
            ["Dockerfile + Compose", "Yes"],
            ["CI/CD workflow", "Yes"],
            ["Editable Word report and PDF report", "Yes"],
            ["Sub-five-minute demonstration video", "Yes"],
        ],
        [4.65, 1.35],
    )
    add_heading(doc, "Limitations and next steps", 2)
    doc.add_paragraph(
        "This is a deliberately lightweight baseline trained on 600 images, not a production pet recognition "
        "model. Accuracy should be improved through transfer learning (for example, MobileNet or ResNet), a "
        "larger stratified dataset, augmentation, calibration, and model fairness checks. The CI/CD "
        "configuration is complete but requires repository secrets, a GHCR namespace, and an external "
        "Compose host to execute the production deployment path."
    )
    add_heading(doc, "References")
    for label, url in (
        ("Kaggle — Dogs vs Cats competition dataset", "https://www.kaggle.com/competitions/dogs-vs-cats/data"),
        ("TensorFlow Datasets — cats_vs_dogs catalog", "https://www.tensorflow.org/datasets/catalog/cats_vs_dogs"),
        ("DVC documentation", "https://dvc.org/doc"),
        ("MLflow tracking documentation", "https://mlflow.org/docs/latest/ml/tracking/"),
        ("FastAPI documentation", "https://fastapi.tiangolo.com/"),
        ("GitHub Container Registry documentation", "https://docs.github.com/packages/working-with-a-github-packages-registry/working-with-the-container-registry"),
        ("Prometheus Python client documentation", "https://prometheus.github.io/client_python/"),
    ):
        add_reference(doc, bullet_id, label, url)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
